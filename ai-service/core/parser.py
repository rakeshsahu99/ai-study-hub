import io
import os
from typing import List, Dict, Any
import pdfplumber
from pypdf import PdfReader, PdfWriter
from docx import Document
from google import genai
from google.genai import types

def extract_text_from_pdf(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF byte stream.
    Tries pdfplumber first, falling back to pypdf on error or empty text.
    """
    pages_data = []
    
    # 1. Try pdfplumber
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        pages_data.append({
                            "text": text.strip(),
                            "metadata": {
                                "source": filename,
                                "page": i + 1
                            }
                        })
                except Exception:
                    pass  # Fall back to page-by-page pypdf below if needed
    except Exception:
        pages_data = []

    # 2. Fallback to pypdf if pdfplumber failed or extracted no text
    if not pages_data:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages_data.append({
                        "text": text.strip(),
                        "metadata": {
                            "source": filename,
                            "page": i + 1
                        }
                    })
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file with all available parsers: {str(e)}")

    return pages_data

def extract_text_from_docx(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Extracts text from a Word document (.docx) byte stream, including paragraphs and tables.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
                
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        combined_text = "\n".join(full_text)
        if combined_text:
            return [{
                "text": combined_text,
                "metadata": {
                    "source": filename,
                    "page": 1
                }
            }]
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")
        
    return []

def extract_text_from_txt(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Extracts text from plain text (.txt) files.
    """
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")  # Fallback decode
        
    if text.strip():
        return [{
            "text": text.strip(),
            "metadata": {
                "source": filename,
                "page": 1
            }
        }]
    return []

# Reuse the existing GenAI client configuration (picks up GEMINI_API_KEY)
api_key = os.getenv("GEMINI_API_KEY")
client = genai.Client(api_key=api_key) if api_key else None

def extract_text_multimodal_pdf(file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF using Gemini 2.5 Flash for layout-aware parsing.
    Falls back to local text extraction on rate limits or API issues.
    """
    if not client:
        print("Gemini client not configured. Falling back to local PDF parser.")
        return extract_text_from_pdf(file_bytes, filename)
        
    pages_data = []
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        total_pages = len(reader.pages)
        
        for i in range(total_pages):
            try:
                # 1. Isolate the target page bytes
                writer = PdfWriter()
                writer.add_page(reader.pages[i])
                page_stream = io.BytesIO()
                writer.write(page_stream)
                page_bytes = page_stream.getvalue()
                
                # 2. Setup structural extraction instruction
                prompt = (
                    "Extract all text, headers, and bullet points from this page. "
                    "If there are tables, represent them as clean Markdown tables. "
                    "If there are mathematical formulas, represent them as standard equations. "
                    "Maintain the structure in clean Markdown format."
                )
                
                # 3. Query Gemini with raw page bytes
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=[
                        types.Part.from_bytes(data=page_bytes, mime_type='application/pdf'),
                        prompt
                    ]
                )
                
                page_text = response.text
                if page_text and page_text.strip():
                    pages_data.append({
                        "text": page_text.strip(),
                        "metadata": {
                            "source": filename,
                            "page": i + 1
                        }
                    })
                else:
                    # Fallback to local page parse if API returns blank
                    local_text = reader.pages[i].extract_text()
                    if local_text and local_text.strip():
                        pages_data.append({
                            "text": local_text.strip(),
                            "metadata": {
                                "source": filename,
                                "page": i + 1
                            }
                        })
            except Exception as page_err:
                print(f"Error parsing page {i+1} with Gemini: {page_err}. Falling back to local page parser.")
                try:
                    local_text = reader.pages[i].extract_text()
                    if local_text and local_text.strip():
                        pages_data.append({
                            "text": local_text.strip(),
                            "metadata": {
                                "source": filename,
                                "page": i + 1
                            }
                        })
                except Exception:
                    pass
    except Exception as e:
        print(f"Failed Gemini multimodal parsing pipeline: {e}. Falling back to local parser.")
        return extract_text_from_pdf(file_bytes, filename)
        
    return pages_data
